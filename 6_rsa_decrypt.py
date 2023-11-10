from docx import Document

def write_to_docx(text, file_path):
    try:
        document = Document(file_path)
    except:
        document = Document()
    
    document.add_paragraph(text)
    document.save(file_path)    

def FiFunc(a, filename):
    b = a
    s = []
    c = 2
    write_to_docx(f'Разложим {a} на простые множители:', filename)
    while c != b + 1:
        if a % c == 0:
            s.append(c)
            a = a / c
            write_to_docx(str(c), filename)
            c = 2
        else:
            c += 1
    i = 0
    for elem in s:
        s[i] -= 1
        i += 1
    i = 0
    a = 1
    for elem in s:
        a *= s[i]
        i += 1
    write_to_docx(f'Fi(n): {a}', filename)
    return a

#возведение в -1 по модулю
def MinusOneMod(exp, fi, filename):
    r = 0
    q = 0
    y  = [0, 1]
    mod = fi
    coef_first = 0
    coef_second = 1
    write_to_docx('(Запишем y в столбце справа:)', filename)
    write_to_docx(f'y[-2] = {y[0]}', filename)
    write_to_docx(f'y[-1] = {y[1]}', filename)

    while r != 1:
        q = fi // exp
        r = fi % exp
        y.append(y[coef_first] - y[coef_second]*q)
        write_to_docx('{} = (q{}){}*{} + r({}){} ,посчитаем y({}) = y({}){} - y({}){}*q({}){} = {}'
              .format(fi, coef_first, q, exp, coef_first, r, coef_first, coef_first - 2, y[coef_first], coef_second - 2,
                    y[coef_second], coef_first, q, y[coef_second + 1]), filename)
        fi = exp
        exp = r
        coef_first += 1
        coef_second += 1
    y[coef_second] %= mod
    write_to_docx(f'Ответ по mod {mod} = {y[coef_second]}', filename)
    return y[coef_second]


def FastExponentiation(num1, num2, num3, filename):
    power = bin(num2)
    write_to_docx(power, filename)
    power = power[2:]
    write_to_docx(f"{num2} в двоичной: {power}", filename)
    r1 = 1
    r2 = num1
    r3 = 1
    for i in range(len(power ) - 1):
        r3 = r1 * (r2 ** int(power[i])) % num3
        xyu = (r2 ** int(power[i]))
        write_to_docx(f"{r1} * {xyu} = {r3}", filename)
        r1 = r3 ** 2 % num3
        write_to_docx(f"{r3} ^2 = {r1}", filename)
    answer = r1 * (r2 ** int(power[len(power ) - 1])) % num3
    xyu1 = r2 ** int(power[len(power ) - 1])
    write_to_docx(f"{r1} * {xyu1} =  {answer}", filename)
    write_to_docx(f'Ответ: {answer}', filename)
    return answer



def rsa_decrypt(n, e, c, filename):
    write_to_docx('Найдем Fi(n):', filename)
    FiN = FiFunc(n, filename)
    write_to_docx('Найдем секретную экспоненту d = e^(-1) mod Fi(n):', filename)
    d = MinusOneMod(e, FiN, filename)
    write_to_docx('Найдем сообщение m = c^d mod n', filename)
    m = FastExponentiation(c, d, n, filename)
    assert(m == c**d %n)
    write_to_docx(f'Cообщение m = {m}', filename)
    write_to_docx('fi(n) = {} Секретная экспонента d = {} сообщение m = {}'.format(FiN, d, m), filename)

