from docx import Document

def write_to_docx(text, file_path):
    try:
        document = Document(file_path)
    except:
        document = Document()
    
    document.add_paragraph(text)
    document.save(file_path)    
#exp^-1 mod fi
def MinusOneMod(exp, fi, filename):
    r = 0
    g = 0
    y  = [0, 1]
    mod = fi
    expAns = exp
    coef_first = 0
    coef_second = 1
    write_to_docx('(Запишем y в столбце справа:)', filename)
    write_to_docx(f'y[-2] = {y[0]}', filename)
    write_to_docx(f'y[-1] = {y[1]}', filename)

    while r != 1:
        g = fi // exp
        r = fi % exp
        y.append(y[coef_first] - y[coef_second]*g)
        write_to_docx('{} = (g{}){}*{} + r({}){} ,посчитаем y({}) = y({}){} - y({}){}*g({}){} = {}'
              .format(fi, coef_first, g, exp, coef_first, r, coef_first, coef_first - 2,
                      y[coef_first], coef_second - 2, y[coef_second], coef_first, g, y[coef_second + 1]), filename)
        fi = exp
        exp = r
        coef_first += 1
        coef_second += 1
    y[coef_second] %= mod
    write_to_docx('Ответ {}^-1 mod {} = {}'.format(expAns, mod, y[coef_second]), filename)
    return y[coef_second]

#подпись
def Sign(p, g, y, k, m, filename):
    write_to_docx('y = g^x mod p:', filename)
    write_to_docx('{} = {}^x mod {}'.format(y, g, p), filename)
    x = 0
    while ((g ** x) - y) % p != 0:
        x += 1
    write_to_docx('Перебором получаем х = {x}', filename)
    a = pow(g, k, p)
    write_to_docx('a = g^k mod p = {}^{} mod {} = {}'.format(g, k, p, a), filename)
    write_to_docx('Посчитаем k^-1 mod p-1:', filename)
    ('Промежуточный расчет: --------')

    b = MinusOneMod(k, p - 1, filename)
    b = pow((m - x * a) * b, 1, p - 1)
    write_to_docx('------------------------------', filename)
    write_to_docx('b = |M - xa|* k^-1 mod(p-1) = ({} - {} * {} ) * {} ^-1 mod {} = {}'.format(m, x, a, k, p - 1, b), filename)
    write_to_docx('Ответ: ({},{})'.format(a, b), filename)

 


#зашифровать
def Encrypt(p, g, y, m, k, filename):
    write_to_docx('y = g^x mod p:', filename)
    write_to_docx('{} = {}^x mod {}'.format(y, g, p), filename)
    x = 0
    while ((g ** x) - y) % p != 0:
        x += 1
    write_to_docx(f'Перебором (или из условия) получаем х = {x}', filename)

    a = pow(g, k, p)
    Temp = pow(y, k, p)
    b = pow(m*Temp, 1, p)

    write_to_docx('a = g^k mod p = {}^{} mod {} = {}'.format(g, k, p, a), filename)
    write_to_docx('b = M * y^k mod p = {} * {} mod {} = {}'.format(m, Temp, p, b), filename)
    write_to_docx('Ответ ({},{})'.format(a, b), filename)

