from docx import Document 

def FastExponentiation(num1, num2, num3, filename):
    power = bin(num2)
    write_to_docx(power, filename)
    power = power[2:]
    write_to_docx(f"{num2} в двоичной: {power}", filename)
    r1 = 1
    r2 = num1
    r3 = 1
    for i in range(len(power ) - 1):
        r3 = r1 * (r2 ** int(power[i])) % num3
        xyu = (r2 ** int(power[i]))
        write_to_docx(f"{r1} * {xyu} = {r3}", filename)
        r1 = r3 ** 2 % num3
        write_to_docx(f"{r3} ^2 = {r1}", filename)
    answer = r1 * (r2 ** int(power[len(power ) - 1])) % num3
    xyu1 = r2 ** int(power[len(power ) - 1])
    write_to_docx(f"{r1} * {xyu1} =  {answer}", filename)
    write_to_docx(f'Ответ: {answer}', filename)
    return answer

def write_to_docx(text, file_path):
    try:
        document = Document(file_path)
    except:
        document = Document()
    
    document.add_paragraph(text)
    document.save(file_path)    


def miller(list_of_numbers, p, filename):
    nums = str(list_of_numbers)
    write_to_docx(f"Проверить, являются ли числа {nums} свидетелями простоты числа {p} по Миллеру. p = {p}", filename)
    write_to_docx(f"Решение:", filename)
    n = p
    for a in list_of_numbers:
        write_to_docx(f"Для числа {a}", filename)
        write_to_docx('Вычислим степени, в которые нужно будет возводить потенциальные свидетели простоты:', filename)
        write_to_docx('Формула для нахождения: (p-1)/ 2^t', filename)
        i = 0
        l = []
        while float(((n - 1)/ 2 ** i) % 2) != 1.0:
            l.append(int((((n - 1)/ 2 ** i))))
            i += 1
        l.append(int((((n - 1)/ 2 ** i))))
        l = l[::-1]
        text = ""
        for i in l:
            text += str(i) + " "
        write_to_docx(text, filename)
        write_to_docx('Воспользуемся быстрым возведением в степень {}^{} mod {}'.format(a, l[0], n), filename)
        expon = FastExponentiation(a, l[0], n, filename) 
        assert(expon == a**(l[0]) % n)

        i = 0
        PS = []
        temp = pow(a, l[0] , n)

        for i,item in enumerate(l):
            write_to_docx('{}^{} mod {} = {}'.format(a, item, n, pow(a, item, n)), filename)
            if i == 0:
                write_to_docx('', filename)
            if i > 0:
                write_to_docx(' = {} ^2 mod {}'.format(temp, n), filename)
            temp = pow(a, item, n)
            PS.append(temp)
            i += 1

        MillerTrigger = False
        NotOne = True

        for i,item in enumerate(PS):
            if item == n - 1 and PS[i + 1] == 1 and PS[len(PS) - 1]:
                MillerTrigger = True
                write_to_docx(f'Т.к. p - 1 элемент находится перед первой единицей, и последний элемент 1, то {a} является свидетелем {n} по Миллеру', filename)
            if item != 1:
                NotOne = False

        if NotOne == True:
            MillerTrigger = True
            write_to_docx('Т.к. все элементы 1, то {} является свидетелем {} по Миллеру'.format(a, n), filename)
        if MillerTrigger == False:
            write_to_docx('Условия простоты по Миллеру не выполнены', filename)
        write_to_docx('---------------------------------------------', filename)

miller([52, 68, 73], 185, "hey3.docx")