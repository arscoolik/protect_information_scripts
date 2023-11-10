from docx import Document 


FILE_NAME = "second_task.docx"

def divisors(n):
    i = 2
    while i * i < n:
        if n % i == 0:
            return i, n // i
        i += 1
    return 1, n

def extended_gcd(a, b):
    if a == 0:
        return b, 0, 1
    else:
        g, x, y = extended_gcd(b % a, a)
        return g, y - (b // a) * x, x

def modinv(a, m):
    g, x, y = extended_gcd(a, m)
    if g != 1:
        raise Exception("Modular inverse does not exist")
    else:
        return x % m

def chinese_remainder_theorem(c1, m1, c2, m2, filename):
    # Solve the system of congruences using CRT
    write_to_docx(f"3.Ищем обратные по модулю x * {m2} = 1 mod {m1} и y * {m1} = 1 mod {m2}", filename)

    # Calculate modular inverses
    inv1 = modinv(m2, m1)
    inv2 = modinv(m1, m2)
    write_to_docx(f"x = {inv1}, y = {inv2}", filename)
    # Calculate the solution using CRT formula
    n = m1 * m2
    write_to_docx(f"4.Из китайской теоремы об остатках, считаем такую штуку :(m * aq * x + e * ap * y) % n" , filename)
    write_to_docx(f"Ответом будет: ({c1} * {m2} * {inv1} + {c2} * {m1} * {inv2})  % {n}", filename)
    first = c1 * m2 * inv1 + c2 * m1 * inv2
    write_to_docx(f"{first} % {n}", filename)
    x = (c1 * m2 * inv1 + c2 * m1 * inv2) % (m1 * m2)
    return x

def pow_with_crt(a, b, p, q, filename):
    # Calculate a^b mod p and a^b mod q using modular exponentiation
    write_to_docx(f"2. На этом этапе нам нужно посчитать {a}^{b} mod {q} и {a}^{b} mod {p}, надеюсь вы сможете сами", filename) 
    ap = pow(a, b, p)
    aq = pow(a, b, q)
    write_to_docx(f"aq = {a}^{b} mod {q} = {aq} и ap = {a}^{b} mod {p} = {ap}", filename)
    # Use CRT to combine the results
    result = chinese_remainder_theorem(ap, p, aq, q, filename) 

    return result

def write_to_docx(text, file_path):
    try:
        document = Document(file_path)
    except:
        document = Document()
    
    document.add_paragraph(text)
    document.save(file_path)

def rsa_encrypt(n, e, m, filename): 
    write_to_docx(f"Условие: Зашифровать сообщение по схеме RSA. Открытый ключ: n = {n}; e = {e}. Сообщение: m = {m}", filename)
    write_to_docx("Зашифрованное сообщение c = m^e mod n", filename)
    write_to_docx("Для этого вопспользуемся китайской теоремой об остатках", filename)
    p, q = divisors(n)
    write_to_docx(f"1. Делители числа n это {p}, {q}", filename)

    ans = pow_with_crt(m, e, p, q, filename)
    write_to_docx(f"Ответ: {ans}", filename)
    return ans
